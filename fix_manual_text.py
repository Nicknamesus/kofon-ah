"""Apply the review's text fixes to the hand-edited root manuals (EN + ZH).

These copies have no builder, so we patch them in place. Fixes:
  - the "Chapter 18" cross-reference (the cookbook is Chapter 20)
  - reword internal-engineering jargon for an external-facing tone
    (hardcoded single-process build / CORS dance / behind the GFW)

Run-aware replace preserves each paragraph's inline formatting.
Run:  python fix_manual_text.py   (with the backend venv)
"""
from __future__ import annotations

import os

from docx import Document

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

EN = os.path.join(_ROOT, "Kofon_Chatbot_Engineering_Manual.docx")
ZH = os.path.join(_ROOT, "Kofon_Chatbot_Engineering_Manual_Chinese.docx")

EN_FIXES = [
    ("hardcoded single-process build", "self-contained single-service build"),
    ("Chapter 18 (Maintenance Cookbook)", "Chapter 20 (Maintenance Cookbook)"),
    ("no separate front-end server, no CORS dance, and no second port",
     "no separate front-end server, no cross-origin configuration to manage, and no second port"),
    ("operates from inside China", "operates from inside mainland China"),
    ("are either blocked or high-latency behind the GFW",
     "are either unavailable or high-latency in that network environment"),
]

ZH_FIXES = [
    ("硬编码单进程构建", "单服务一体化构建"),
    ("第 18 章", "第 20 章"),
    ("没有 CORS 折腾", "无需处理跨域（CORS）配置"),
    ("在 GFW 之后要么被封锁，要么高延迟", "在中国大陆网络环境下要么无法访问，要么高延迟"),
]


def replace_in_paragraph(p, old: str, new: str) -> int:
    """Replace `old` with `new` across runs, keeping the start run's format."""
    runs = p.runs
    text = "".join(r.text for r in runs)
    count = 0
    while True:
        start = text.find(old)
        if start < 0:
            break
        end = start + len(old)
        cum, bounds = 0, []
        for r in runs:
            bounds.append((cum, cum + len(r.text), r))
            cum += len(r.text)
        for s, e, r in bounds:
            if e <= start or s >= end:
                continue
            ls, le = max(start, s) - s, min(end, e) - s
            before, after = r.text[:ls], r.text[le:]
            r.text = (before + new + after) if (s <= start < e) else (before + after)
        count += 1
        text = "".join(r.text for r in p.runs)
        runs = p.runs
    return count


def apply(path: str, fixes) -> None:
    doc = Document(path)
    paras = list(doc.paragraphs)
    for tbl in doc.tables:  # include table cells, just in case
        for row in tbl.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    totals = {old: 0 for old, _ in fixes}
    for p in paras:
        for old, new in fixes:
            totals[old] += replace_in_paragraph(p, old, new)
    doc.save(path)
    print(f"== {os.path.basename(path)} ==")
    for i, (old, _) in enumerate(fixes):
        n = totals[old]
        flag = "" if n else "   <-- NOT FOUND"
        print(f"  fix #{i + 1}: {n}x{flag}")


def main():
    apply(EN, EN_FIXES)
    apply(ZH, ZH_FIXES)


if __name__ == "__main__":
    main()
