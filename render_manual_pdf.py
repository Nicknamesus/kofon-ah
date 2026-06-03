"""Render the engineering manual .docx files to polished landscape-A5 PDFs.

Word's COM PDF export (ExportAsFixedFormat) hangs in this headless session, so
instead of converting we re-render: parse the .docx with python-docx and lay it
out with reportlab. This gives a real computed Table of Contents (correct page
numbers + clickable PDF bookmarks), landscape-A5 geometry, running headers /
footers, and CJK font support for the Chinese edition.

Run:  python render_manual_pdf.py        (with the ai-agent-backend venv)
Output: *_A5.pdf next to each source .docx.
"""
from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from reportlab.lib.colors import HexColor, white
from reportlab.lib.pagesizes import A5, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph as RLParagraph,
    Preformatted,
    Spacer,
    Table as RLTable,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents

_HERE = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.dirname(_HERE)

NAVY = HexColor("#132178")
GREY = HexColor("#5A6473")
CODE_BG = HexColor("#F4F5F9")
CODE_FG = HexColor("#1B1B1B")
MAGENTA = HexColor("#8A1C5C")
GRID = HexColor("#D2D6E2")
ZEBRA = HexColor("#F4F5F9")
FILEBOX_BG = HexColor("#EEF0F8")

PAGE = landscape(A5)            # 210 x 148 mm
ML, MR = 16 * mm, 16 * mm
MT, MB = 17 * mm, 15 * mm
USABLE = PAGE[0] - ML - MR

COVER = {
    False: dict(
        title="Kofon Chatbot",
        subtitle="Engineering & Maintenance Manual",
        tagline="The self-contained single-service build  ·  widget · FastAPI · LangGraph · DeepSeek",
        desc=["A guide for engineers who maintain, extend, and operate the assistant.",
              "Walks every file and folder, plus the concepts that tie them together."],
        running="Kofon Chatbot  ·  Engineering & Maintenance Manual",
        toc="Contents",
    ),
    True: dict(
        title="Kofon Chatbot",
        subtitle="工程与维护手册",
        tagline="单服务一体化构建  ·  widget · FastAPI · LangGraph · DeepSeek",
        desc=["面向维护、扩展和运营该助手的工程师的指南。",
              "逐一讲解每个文件和文件夹，以及将它们联系在一起的概念。"],
        running="Kofon Chatbot  ·  工程与维护手册",
        toc="目录",
    ),
}

SOURCES = [
    (os.path.join(_ROOT, "Kofon_Chatbot_Engineering_Manual.docx"), False),
    (os.path.join(_ROOT, "Kofon_Chatbot_Engineering_Manual_Chinese.docx"), True),
]

_BOXDRAW = {
    "│": "|", "─": "-", "├": "|", "└": "\\",
    "┌": "+", "┐": "+", "┘": "+", "┤": "|",
    "┬": "+", "┴": "+", "┼": "+",
}


# --------------------------------------------------------------------------
# docx helpers
# --------------------------------------------------------------------------
def _iter_block_items(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _is_toc_field(p) -> bool:
    return any("TOC" in (t.text or "") for t in p._p.iter(qn("w:instrText")))


def _esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _is_mono(run) -> bool:
    if (run.font.name or "") in ("Consolas", "Courier New", "Courier"):
        return True
    col = run.font.color
    return bool(col is not None and col.rgb is not None and str(col.rgb) == "8A1C5C")


def _runs_markup(p, bold_font: str) -> str:
    out = []
    for run in p.runs:
        txt = _esc(run.text or "")
        if not txt:
            continue
        if _is_mono(run):
            out.append(f'<font face="Courier" color="#8A1C5C">{txt}</font>')
        elif run.bold or run.font.bold:
            out.append(f'<font face="{bold_font}">{txt}</font>')
        else:
            out.append(txt)
    return "".join(out)


def _cell_markup(cell, bold_font: str) -> str:
    return "<br/>".join(_runs_markup(p, bold_font) or _esc(p.text) for p in cell.paragraphs)


# --------------------------------------------------------------------------
# document template: TOC wiring + running header/footer
# --------------------------------------------------------------------------
class ManualDoc(BaseDocTemplate):
    chapter = ""
    running = ""
    hf_font = "Helvetica"

    def afterFlowable(self, flowable):
        bm = getattr(flowable, "_bookmark", None)
        if bm is None:
            return
        level = flowable._toc_level
        text = flowable.getPlainText()
        if level == 0:
            self.chapter = text
        self.canv.bookmarkPage(bm)
        self.canv.addOutlineEntry(text, bm, level=level, closed=(level > 0))
        if level <= 1:
            self.notify("TOCEntry", (level, text, self.page, bm))

    def decorate(self, canvas, doc):
        if doc.page == 1:          # clean cover
            return
        canvas.saveState()
        # header
        canvas.setFont(self.hf_font, 7.5)
        canvas.setFillColor(GREY)
        canvas.drawString(ML, PAGE[1] - 11 * mm, self.running)
        if self.chapter:
            canvas.drawRightString(PAGE[0] - MR, PAGE[1] - 11 * mm, self.chapter)
        canvas.setStrokeColor(GRID)
        canvas.setLineWidth(0.5)
        canvas.line(ML, PAGE[1] - 12.3 * mm, PAGE[0] - MR, PAGE[1] - 12.3 * mm)
        # footer
        canvas.line(ML, MB - 3 * mm, PAGE[0] - MR, MB - 3 * mm)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(GREY)
        canvas.drawCentredString(PAGE[0] / 2.0, MB - 6.5 * mm, str(doc.page))
        canvas.restoreState()


def _styles(cjk: bool):
    base = "STSong-Light" if cjk else "Helvetica"
    bold = "STSong-Light" if cjk else "Helvetica-Bold"
    ss = getSampleStyleSheet()
    s = {"base": base, "bold": bold}
    s["body"] = ParagraphStyle("body", parent=ss["Normal"], fontName=base,
                               fontSize=9.5, leading=13.5, spaceAfter=5)
    s["h1"] = ParagraphStyle("h1", fontName=bold, fontSize=18, leading=22,
                             textColor=NAVY, spaceBefore=0, spaceAfter=2)
    s["h2"] = ParagraphStyle("h2", fontName=bold, fontSize=12.5, leading=16,
                             textColor=NAVY, spaceBefore=9, spaceAfter=3)
    s["h3"] = ParagraphStyle("h3", fontName=bold, fontSize=11, leading=14,
                             textColor=HexColor("#33406B"), spaceBefore=7, spaceAfter=2)
    s["bullet"] = ParagraphStyle("bullet", parent=s["body"], leftIndent=14,
                                 bulletIndent=2, spaceAfter=3)
    s["code"] = ParagraphStyle("code", fontName="Courier", fontSize=8,
                               leading=10.2, textColor=CODE_FG)
    s["filebox"] = ParagraphStyle("filebox", fontName="Courier", fontSize=8,
                                  leading=11, textColor=NAVY)
    s["cell"] = ParagraphStyle("cell", parent=s["body"], fontSize=8.5,
                               leading=11, spaceAfter=0)
    s["cellh"] = ParagraphStyle("cellh", parent=s["cell"], fontName=bold,
                                textColor=white)
    s["cover_t"] = ParagraphStyle("cover_t", fontName=bold, fontSize=34,
                                  leading=40, textColor=NAVY, alignment=1)
    s["cover_s"] = ParagraphStyle("cover_s", fontName=base, fontSize=17,
                                  leading=22, textColor=GREY, alignment=1)
    s["cover_tag"] = ParagraphStyle("cover_tag", fontName=base, fontSize=10.5,
                                    leading=15, textColor=GREY, alignment=1)
    s["cover_d"] = ParagraphStyle("cover_d", fontName=base, fontSize=10,
                                  leading=15, textColor=GREY, alignment=1)
    s["toc_h"] = ParagraphStyle("toc_h", fontName=bold, fontSize=18,
                                leading=22, textColor=NAVY, spaceAfter=4)
    return s


def _toc_flowable(cjk: bool):
    base = "STSong-Light" if cjk else "Helvetica"
    bold = "STSong-Light" if cjk else "Helvetica-Bold"
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("toc1", fontName=bold, fontSize=10.5, leading=17,
                       textColor=NAVY, spaceBefore=5),
        ParagraphStyle("toc2", fontName=base, fontSize=9, leading=13.5,
                       leftIndent=16, textColor=CODE_FG),
    ]
    return toc


# --------------------------------------------------------------------------
# table rendering
# --------------------------------------------------------------------------
def _is_code_table(table) -> bool:
    return len(table.columns) == 1 and len(table.rows) == 1


def _render_code_table(table, st):
    lines = [para.text for para in table.rows[0].cells[0].paragraphs]
    text = "\n".join(lines).strip("\n") or " "
    for k, v in _BOXDRAW.items():
        text = text.replace(k, v)
    pre = Preformatted(text, st["code"])
    bar = 3
    t = RLTable([["", pre]], colWidths=[bar, USABLE - bar])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), NAVY),
        ("BACKGROUND", (1, 0), (1, 0), CODE_BG),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (0, 0), 0), ("RIGHTPADDING", (0, 0), (0, 0), 0),
        ("LEFTPADDING", (1, 0), (1, 0), 9), ("RIGHTPADDING", (1, 0), (1, 0), 9),
        ("TOPPADDING", (1, 0), (1, 0), 6), ("BOTTOMPADDING", (1, 0), (1, 0), 6),
    ]))
    return t


def _render_grid_table(table, st):
    ncols = len(table.columns)
    data = []
    for ri, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            if ri == 0:
                cells.append(RLParagraph(_esc(cell.text) or " ", st["cellh"]))
            else:
                cells.append(RLParagraph(_cell_markup(cell, st["bold"]) or " ", st["cell"]))
        data.append(cells)
    if ncols == 2:
        widths = [USABLE * 0.32, USABLE * 0.68]
    elif ncols == 3:
        widths = [USABLE * 0.24, USABLE * 0.20, USABLE * 0.56]
    else:
        widths = [USABLE / ncols] * ncols
    t = RLTable(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, ZEBRA]),
        ("LINEBELOW", (0, 0), (-1, -1), 0.4, GRID),
        ("LINEAFTER", (0, 0), (-2, -1), 0.4, GRID),
        ("BOX", (0, 0), (-1, -1), 0.5, GRID),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def _render_filebox(markup, st):
    t = RLTable([[RLParagraph(markup, st["filebox"])]], colWidths=[USABLE])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), FILEBOX_BG),
        ("LINEBEFORE", (0, 0), (-1, -1), 2.5, NAVY),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


# --------------------------------------------------------------------------
# build the story
# --------------------------------------------------------------------------
def build_story(doc, st, cjk: bool):
    cov = COVER[cjk]
    blocks = list(_iter_block_items(doc))
    first_h1 = next((i for i, b in enumerate(blocks)
                     if isinstance(b, Paragraph) and b.style and b.style.name == "Heading 1"), 0)

    story = []

    # ---- cover ----
    story.append(Spacer(1, 96))
    story.append(RLParagraph(cov["title"], st["cover_t"]))
    story.append(Spacer(1, 6))
    story.append(RLParagraph(cov["subtitle"], st["cover_s"]))
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="36%", thickness=2, color=NAVY,
                            spaceBefore=2, spaceAfter=18, hAlign="CENTER"))
    story.append(RLParagraph(cov["tagline"], st["cover_tag"]))
    story.append(Spacer(1, 40))
    for line in cov["desc"]:
        story.append(RLParagraph(line, st["cover_d"]))
    story.append(PageBreak())

    # ---- table of contents ----
    story.append(RLParagraph(cov["toc"], st["toc_h"]))
    story.append(HRFlowable(width="100%", thickness=1, color=GRID,
                            spaceBefore=2, spaceAfter=10))
    story.append(_toc_flowable(cjk))
    story.append(PageBreak())

    # ---- body ----
    bm_seq = [0]
    num = [0]

    def bookmark():
        bm_seq[0] += 1
        return f"bm{bm_seq[0]}"

    first_chapter = [True]

    for b in blocks[first_h1:]:
        if isinstance(b, Table):
            num[0] = 0
            story.append(Spacer(1, 2))
            story.append(_render_code_table(b, st) if _is_code_table(b)
                         else _render_grid_table(b, st))
            story.append(Spacer(1, 5))
            continue

        p = b
        if _is_toc_field(p):
            continue
        name = p.style.name if p.style else "Normal"
        text = p.text.strip()

        if name == "Heading 1":
            num[0] = 0
            if not first_chapter[0]:
                story.append(PageBreak())
            first_chapter[0] = False
            head = RLParagraph(_runs_markup(p, st["bold"]) or _esc(text), st["h1"])
            head._toc_level = 0
            head._bookmark = bookmark()
            story.append(KeepTogether([
                head,
                HRFlowable(width="100%", thickness=1.5, color=NAVY,
                           spaceBefore=3, spaceAfter=9),
            ]))
            continue

        if name == "Heading 2":
            num[0] = 0
            head = RLParagraph(_runs_markup(p, st["bold"]) or _esc(text), st["h2"])
            head._toc_level = 1
            head._bookmark = bookmark()
            story.append(head)
            continue

        if name == "Heading 3":
            num[0] = 0
            story.append(RLParagraph(_runs_markup(p, st["bold"]) or _esc(text), st["h3"]))
            continue

        if not text or text in ("Contents", "目录"):
            continue

        markup = _runs_markup(p, st["bold"]) or _esc(text)

        if text.startswith("■"):
            num[0] = 0
            story.append(_render_filebox(markup.lstrip("■ ").lstrip(), st))
        elif name == "List Bullet":
            num[0] = 0
            story.append(RLParagraph(markup, st["bullet"], bulletText="•"))
        elif name == "List Number":
            num[0] += 1
            story.append(RLParagraph(markup, st["bullet"], bulletText=f"{num[0]}."))
        else:
            num[0] = 0
            story.append(RLParagraph(markup, st["body"]))

    return story


def render(src: str, cjk: bool):
    if not os.path.exists(src):
        print(f"skip (not found): {src}")
        return
    if cjk:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    doc = Document(src)
    st = _styles(cjk)
    out = src.replace(".docx", "_A5.pdf")

    pdf = ManualDoc(out, pagesize=PAGE, leftMargin=ML, rightMargin=MR,
                    topMargin=MT, bottomMargin=MB, title=os.path.basename(out))
    pdf.running = COVER[cjk]["running"]
    pdf.hf_font = st["base"]
    frame = Frame(ML, MB, USABLE, PAGE[1] - MT - MB, id="main",
                  leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    # decorate at page END so the running header reflects the chapter that
    # actually starts on this page (the H1 is processed before onPageEnd)
    pdf.addPageTemplates([PageTemplate(id="all", frames=[frame], onPageEnd=pdf.decorate)])

    pdf.multiBuild(build_story(doc, st, cjk))
    print(f"wrote {out}  ({pdf.page} pages)")


def main():
    for src, cjk in SOURCES:
        render(src, cjk)


if __name__ == "__main__":
    main()
