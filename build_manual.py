"""Generate the Kofon Chatbot engineering manual as a .docx.

Run:  python build_manual.py
Output: Kofon_Chatbot_Engineering_Manual.docx
"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor, Inches

from toc_util import add_toc_before_first_heading

NAVY = RGBColor(0x13, 0x21, 0x78)
GREY = RGBColor(0x5A, 0x64, 0x73)
CODE_BG = "F2F3F7"

doc = Document()

# ---------------- base styles ----------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for i, sz in ((1, 20), (2, 15), (3, 12.5)):
    st = doc.styles[f"Heading {i}"]
    st.font.name = "Calibri"
    st.font.color.rgb = NAVY
    st.font.size = Pt(sz)
    st.font.bold = True


def add_page_numbers(document):
    """Add a centered 'Page N' field to every section footer."""
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # clear any existing content
        for r in list(para.runs):
            r._element.getparent().remove(r._element)

        def _field(run, instr):
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr_el = OxmlElement("w:instrText")
            instr_el.set(qn("xml:space"), "preserve")
            instr_el.text = instr
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run._r.append(fld_begin)
            run._r.append(instr_el)
            run._r.append(fld_end)

        pre = para.add_run("Page ")
        pre.font.size = Pt(9)
        pre.font.color.rgb = GREY
        num = para.add_run()
        num.font.size = Pt(9)
        num.font.color.rgb = GREY
        _field(num, " PAGE ")


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def H1(text):
    doc.add_paragraph().add_run().add_break()  # spacing
    h = doc.add_heading(text, level=1)
    return h


def H2(text):
    return doc.add_heading(text, level=2)


def H3(text):
    return doc.add_heading(text, level=3)


def P(text="", style=None):
    p = doc.add_paragraph(style=style)
    _emit_runs(p, text)
    return p


def _emit_runs(p, text):
    """Minimal inline markup: **bold**, `code`."""
    import re

    tokens = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x8A, 0x1C, 0x5C)
        else:
            p.add_run(tok)


def BULLET(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _emit_runs(p, it)


def NUM(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        _emit_runs(p, it)


def CODE(text):
    text = text.strip("\n")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, CODE_BG)
    cell.width = Inches(6.3)
    cell.paragraphs[0].text = ""
    first = True
    for line in text.split("\n"):
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1B)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def KV(rows, headers=("Field", "Meaning")):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].add_run(htext).bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            _emit_runs(cells[i].paragraphs[0], str(val))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def PAGEBREAK():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def FILEBOX(label):
    """Small monospace caption naming the file(s) a section documents."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("■  " + label)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = NAVY
    r.bold = True


# ============================================================
# COVER
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(4):
    title.add_run().add_break()
r = title.add_run("Kofon Chatbot")
r.font.size = Pt(34)
r.font.bold = True
r.font.color.rgb = NAVY
title.add_run().add_break()
r = title.add_run("Engineering & Maintenance Manual")
r.font.size = Pt(20)
r.font.color.rgb = GREY
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run().add_break()
r = sub.add_run("The self-contained single-service build (widget + FastAPI + LangGraph + DeepSeek)")
r.italic = True
r.font.color.rgb = GREY
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(8):
    meta.add_run().add_break()
r = meta.add_run("A guide for engineers who maintain, extend, and operate the bot.\n")
r.font.color.rgb = GREY
r = meta.add_run("Walks every file and folder, plus the concepts that tie them together.")
r.font.color.rgb = GREY
PAGEBREAK()

# ============================================================
# 0. HOW TO READ THIS MANUAL
# ============================================================
doc.add_heading("How to read this manual", level=1)
P(
    "This manual is organised by the layout of the repository. Each chapter covers one "
    "module, folder, or layer, names the file(s) it documents in a blue monospace banner, "
    "and explains both the mechanics (what the code does) and the intent (why it is built "
    "that way). Repetitive files — the dozens of per-family seed YAMLs, the four migration "
    "files, the per-language translation columns — are explained once per **type**, with a "
    "single representative example, rather than file by file."
)
P(
    "If you are new to the project, read Chapter 1 (Architecture) first; it is the map that "
    "makes every later chapter make sense. If you are here to perform a specific task — add a "
    "product, add a language, swap the CRM — jump to Chapter 20 (Maintenance Cookbook), which "
    "cross-references the detailed chapters."
)
P("Conventions used throughout:")
BULLET([
    "`monospace` marks a filename, symbol, env var, or literal value.",
    "A grey box is a verbatim code or shell excerpt.",
    "“Phase 1–4” refer to the build's historical milestones; the comments in the "
    "code still use them, so the manual keeps the vocabulary. Phase 1 = data + read tools, "
    "Phase 2 = the agent graph, Phase 3 = embeddings + configurator, Phase 4 = CRM/email side "
    "effects.",
])

# ============================================================
# 1. ARCHITECTURE
# ============================================================
H1("1. Architecture overview")
P(
    "The Kofon chatbot is a B2B pre-sales / guide / post-sales assistant for a motion-components "
    "manufacturer. This repository is the **self-contained single-service build**: one FastAPI "
    "process serves both the browser widget (static files) and the JSON/SSE API from the same "
    "origin. There is no separate front-end server, no cross-origin configuration to manage, and "
    "no second port. A "
    "decoupled, reusable build of the same system lives in sibling repositories "
    "(`../ai-agent-backend` + `../ai-agent-addon`); this one trades that modularity for the "
    "simplest possible deployment."
)

H2("The stack")
KV([
    ["Browser widget", "Vanilla JS (no framework). Renders chat + structured “cards”. In `widget/`."],
    ["Web framework", "FastAPI + Uvicorn. Entry point `app/main.py`, served via `app/serve.py`."],
    ["Agent engine", "LangGraph state machine. The brain lives in `app/agent/`."],
    ["LLM", "DeepSeek (`deepseek-chat` / `deepseek-reasoner`) via `langchain-deepseek`."],
    ["Database", "PostgreSQL 16 with the pgvector extension (`pgvector/pgvector:pg16`)."],
    ["ORM / migrations", "SQLAlchemy 2 (async, asyncpg) + Alembic."],
    ["Embeddings", "Pluggable: hash (dev), BGE-M3 (local), or DashScope (Alibaba)."],
    ["CRM / email", "Pluggable adapters: log (default), Zoho CRM, Aliyun DirectMail."],
], headers=("Layer", "Technology"))

H2("Why DeepSeek and Chinese-native services")
P(
    "Kofon operates from inside mainland China, so the usual Western SaaS defaults (OpenAI, "
    "Anthropic, SendGrid, Salesforce) are either unavailable or high-latency in that network "
    "environment. Every external "
    "dependency therefore has a China-friendly option: DeepSeek for the LLM, DashScope "
    "(Alibaba Qwen) for embeddings, Zoho's `.com.cn` region for CRM, and Aliyun DirectMail for "
    "email. This single constraint explains many design choices you will meet later; keep it "
    "in mind."
)

H2("The request lifecycle (the one diagram to memorise)")
P(
    "Every user turn — typed text or a button click — is a single POST to `/api/messages` that "
    "streams Server-Sent Events back. End to end:"
)
NUM([
    "The widget (`widget/api.js`) POSTs `{session_uuid, text|gate_choice|…}` to `/api/messages`.",
    "The SSE router (`app/routers/agent.py`) sanitises the input, writes analytics rows "
    "(`app/persistence.py`), then opens a LangGraph checkpointer and builds the graph.",
    "The graph (`app/agent/graph.py`) dispatches to one flow node based on saved state. "
    "Nodes call DeepSeek and the read-only tools, and return partial state updates "
    "(messages, cards, slots, outcome).",
    "The router streams each update out as `bot_text` / `card` / `outcome` events and persists "
    "them. When the graph finishes it sends `done`.",
    "If the turn reached a terminal outcome (sell / human / resolved), the outcome node already "
    "fired the side effects (CRM record + division email) through `app/sideeffects/`.",
    "The widget renders text bubbles and cards as the events arrive.",
])
P(
    "State survives between turns in two places: the **LangGraph checkpoint** (full graph state, "
    "keyed by `session_uuid`, used to resume) and the **`conversations`/`messages` tables** (a "
    "slim, queryable mirror for analytics). Understanding that split resolves most "
    "“where is this stored?” questions."
)

H2("Repository map")
KV([
    ["`app/`", "All backend Python. The subsections below own one chapter each."],
    ["`app/agent/`", "The LangGraph graph, state, nodes, LLM factory, checkpointer, safety."],
    ["`app/tools/`", "Read-only catalog/KB query functions the nodes call."],
    ["`app/models/`", "SQLAlchemy ORM tables (content, runtime, side-effects)."],
    ["`app/routers/`", "FastAPI endpoints (the SSE agent route + thin tool routes)."],
    ["`app/sideeffects/`", "Outcome orchestration: CRM payloads, routing, audit."],
    ["`app/crm/`, `app/mail/`", "Pluggable provider adapters."],
    ["`app/seed/`", "Loaders and catalog-ingest scripts (file → Postgres)."],
    ["`seed/`", "The editable content itself: YAML + CSV."],
    ["`alembic/`", "Database migrations."],
    ["`widget/`", "The browser chat widget and demo page."],
    ["root", "Deployment + project config: `deploy.sh`, `docker-compose.yml`, `.env.example`, `pyproject.toml`."],
], headers=("Path", "Role"))
PAGEBREAK()

# ============================================================
# 2. CONFIG & RUNTIME ENTRY
# ============================================================
H1("2. Configuration, runtime, and the FastAPI entry point")
FILEBOX("app/config.py  ·  app/runtime.py  ·  app/serve.py  ·  app/main.py  ·  app/db.py")

H2("app/config.py — one settings object, loaded once")
P(
    "`Settings` is a pydantic-settings `BaseSettings`. It reads `.env` (and real environment "
    "variables, which win) and exposes every knob the app has: Postgres parts, the DeepSeek "
    "keys/models, the embedding provider, and the Phase-4 CRM/email/routing settings. "
    "`get_settings()` is `@lru_cache`d, so the object is constructed once per process and "
    "imported everywhere."
)
P(
    "Two computed properties matter. `effective_database_url` composes the asyncpg URL used by "
    "the app's SQLAlchemy engine (or returns `DATABASE_URL` verbatim if you set it). "
    "`checkpointer_database_url` rewrites that same URL to the plain `postgresql://` (psycopg) "
    "scheme, because the LangGraph checkpointer ships on psycopg, not asyncpg. **They point at "
    "the same database with two different drivers** — a recurring theme."
)
P(
    "Maintenance note: to add a setting, add a typed field with a sane default and document it "
    "in `.env.example`. `extra=\"ignore\"` means stray env vars are silently dropped, so a typo "
    "in `.env` fails quietly — double-check spelling."
)

H2("app/runtime.py & app/serve.py — the Windows event-loop fix")
P(
    "psycopg's async driver cannot run on Windows' default `ProactorEventLoop`; it needs the "
    "`SelectorEventLoop`. `install_async_event_loop_policy()` swaps the policy and is a no-op on "
    "Linux/macOS. It must be called **before** any event loop is created — which is why "
    "`app/main.py` calls it at import time."
)
P(
    "On Python 3.14 even that isn't enough, because the loop policy no longer reliably governs "
    "the loop uvicorn ends up on. `app/serve.py` is the robust answer: it constructs a "
    "`SelectorEventLoop` itself and runs uvicorn's `Server.serve()` inside it. **Always start "
    "the server with `python -m app.serve --port 8001`, not bare `uvicorn app.main:app`** — on "
    "Windows the latter will crash the checkpointer. On Linux the two are equivalent."
)

H2("app/main.py — the ASGI app")
P("`main.py` wires the FastAPI application together:")
BULLET([
    "A `lifespan` context pings the DB (`SELECT 1`) on startup and disposes the engine on "
    "shutdown — a fail-fast check that Postgres is reachable.",
    "CORS is wide-open (`allow_origins=[\"*\"]`) purely as a safety net for third-party "
    "embeds; same-origin serving means the browser never even sends a preflight.",
    "The tool and agent routers are mounted, then `/api/health` is defined.",
    "Crucially, the static `widget/` mount at `/` comes **last**. FastAPI matches explicit "
    "routes first and falls through to the mount, so every `/api/*` route must be registered "
    "before the catch-all or it would be shadowed by the static server.",
])

H2("app/db.py — engine, session factory, Base")
P(
    "Defines the async SQLAlchemy `engine` (with `pool_pre_ping=True` so stale pooled "
    "connections are detected), the `SessionLocal` async-session factory used throughout the "
    "nodes and tools, the declarative `Base` every model inherits, and the `get_session()` "
    "FastAPI dependency. Nodes open short-lived sessions with `async with SessionLocal() as "
    "session:` and release them before slow LLM calls so a pool connection isn't held during "
    "the network hop."
)
PAGEBREAK()

# ============================================================
# 3. THE AGENT GRAPH
# ============================================================
H1("3. The agent graph and shared state")
FILEBOX("app/agent/graph.py  ·  app/agent/state.py")
P(
    "This is the heart of the system. If you read one chapter closely, read this one — every "
    "node chapter assumes it."
)

H2("The state schema (state.py)")
P(
    "`AgentState` is a `TypedDict(total=False)`: a flat bag of channels the whole graph shares. "
    "LangGraph runs each node with the full state and merges back whatever partial dict the "
    "node returns. The important channels:"
)
KV([
    ["`session_uuid` / `conversation_id`", "Stable identifiers injected by the router."],
    ["`flow`", "Which of the four primary flows owns the turn: presales | guide | postsales | other."],
    ["`current_node`", "Diagnostic breadcrumb, e.g. `guide.find.presented`."],
    ["`slots`", "Free-form collected facts. The reducer `merge_slots` shallow-merges them."],
    ["`messages`", "The chat log; the `add_messages` reducer **appends** rather than replaces."],
    ["`cards`", "Structured UI events to stream; appended each turn."],
    ["`outcome`", "Terminal marker, set only by outcome_* nodes: sell | human_handoff | resolved."],
    ["`language`", "The widget-selected language code (EN/DE/…); every node honours it."],
], headers=("Channel", "Purpose"))
P(
    "The reducers are the subtle part. `messages` and `cards` accumulate; `slots` merges "
    "shallowly, which means **a node that owns a nested slot dict must return the whole dict, "
    "not a partial** — `merge_slots` replaces nested dicts wholesale. Most node bugs trace back "
    "to forgetting this. Slots are namespaced per flow (`slots.presales.*`, `slots.postsales.*`, "
    "`slots.filters`, `slots.candidates`, …) so flows don't collide."
)

H2("How the graph is built (graph.py)")
P(
    "`build_graph(checkpointer)` registers every node and wires conditional edges. There is one "
    "real entry decision and a small set of “after-node” routers. The design rule: a node "
    "either **ends the turn** (returns to the user and waits) or **hands off in the same turn** "
    "to the next node. Conditional-edge functions read the post-node state and choose."
)
P("The entry dispatch (`_entry_dispatch`) runs on every turn and decides where to start:")
BULLET([
    "If `outcome` is already set, the flow has terminated; route to `post_outcome_chat` so the "
    "user can keep talking without re-running the flow.",
    "If `slots.force_human` is set (the “Talk to a human” utility chip), jump straight to "
    "`outcome_human`.",
    "If no `flow` is set yet, run `entry_router` (the LLM classifier).",
    "Otherwise dispatch into the flow: presales → `presales.figure_out`; guide → the "
    "`_guide_dispatch` sub-decision; postsales → `_postsales_dispatch`; other → "
    "`other.reclassify`.",
])
P(
    "`_guide_dispatch` and `_postsales_dispatch` are pure functions of the slot state. For "
    "example, guide routes to `guide.customize` when a customise sub-flow is active, to "
    "`guide.happy_gate` when results were already presented and the user hasn't answered, and "
    "to `guide.find` otherwise. Postsales routes to `match_kb` when a problem was picked or a "
    "symptom is ready, to `fix_gate` after a solution was presented, and back to `identify` "
    "when an ambiguous shortlist needs refining."
)
P(
    "The graph shape in one glance (reproduced from the module docstring):"
)
CODE(
    "START\n"
    "  ├─ (flow set)  → flow node\n"
    "  └─ (no flow)   → entry_router → flow node\n\n"
    "presales.figure_out → guide.find (handoff) | END\n"
    "guide.find / guide.customize → END (await reply, then happy_gate)\n"
    "guide.happy_gate → outcome_sell | outcome_human | END\n"
    "postsales.identify → postsales.match_kb | END\n"
    "postsales.match_kb → END (await reply, then fix_gate) | outcome_human\n"
    "postsales.fix_gate → outcome_resolved | outcome_human | match_kb (retry) | END\n"
    "other.reclassify → re-dispatch via flow | END\n"
    "outcome_* → END"
)
P(
    "`build_echo_graph` is a one-node smoke graph kept around to prove the LangGraph + DeepSeek "
    "+ checkpointer plumbing works in isolation. It is not used in production."
)
P(
    "**To add a flow node:** write it under `app/agent/nodes/`, register it in `build_graph`, "
    "and extend the relevant dispatch function plus the destination list of every conditional "
    "edge that can reach it (LangGraph validates the allowed targets, so a missing entry raises "
    "at build time rather than silently dropping the edge)."
)
PAGEBREAK()

# ============================================================
# 4. LLM FACTORY, SAFETY, I18N
# ============================================================
H1("4. The LLM factory, prompt safety, and i18n")
FILEBOX("app/agent/llm.py  ·  app/agent/sanitize.py  ·  app/i18n.py")

H2("app/agent/llm.py — every node's model comes from here")
P(
    "Centralising model construction means switching models (per node or globally) is a "
    "one-line change. `get_chat_llm(temperature=, model=)` returns a `deepseek-chat` model — "
    "fast and cheap, supports tool calls and structured output — and is what almost every node "
    "uses. `get_reasoner_llm()` returns the slower `deepseek-reasoner` for harder reasoning; "
    "it is reserved, not used by default."
)
P(
    "The key helper is `system_message(content, lang)`. It builds a `SystemMessage` and appends "
    "two things to every prompt: the **language instruction** (so DeepSeek replies in the "
    "user's language) and `PROMPT_ARMOR`, a fixed paragraph telling the model that user text is "
    "data to analyse, never instructions to obey. Use `system_message(…)` for every node "
    "prompt; do not hand-roll a bare `SystemMessage`, or you lose both protections."
)

H2("app/agent/sanitize.py — input hygiene and content fencing")
P("Two small, important functions:")
BULLET([
    "`sanitize_user_input(text)` strips control characters, trims, and truncates to 2000 "
    "chars. The SSE router runs every inbound message through it before it reaches the graph. "
    "It returns `None` for empty input.",
    "`fence(value, label)` wraps any externally-sourced string in XML-style tags so the LLM "
    "treats it as data. Nodes that interpolate DB rows or JSON into a prompt (the gate nodes' "
    "“answer a question about the candidates” branch) fence that content first. This is "
    "defence-in-depth alongside `PROMPT_ARMOR`.",
])

H2("app/i18n.py — translating everything the bot says")
P(
    "The widget header lets the user pick one of seven languages (EN, DE, FR, RU, JA, KO, ZH). "
    "Two halves of the problem are solved here:"
)
BULLET([
    "**Hardcoded strings.** `t(key, lang, **kwargs)` looks up a translation in the big `_T` "
    "table and `.format()`s placeholders. A missing language falls back to EN rather than "
    "throwing; a missing **key** raises (the table is the contract).",
    "**LLM-generated strings.** `language_instruction(lang)` returns the sentence appended to "
    "every system prompt instructing the model to write in that language while keeping SKUs and "
    "units untranslated.",
])
P(
    "`_T` is a dict of `key → {lang: string}`. Every user-facing sentence the backend emits "
    "has an entry. **To add a language:** add it to `SUPPORTED_LANGUAGES` and `LANGUAGE_NAMES`, "
    "then add that column to every entry in `_T` (missing columns fall back to EN, so you can "
    "ship incrementally). Remember the widget has its own parallel translation table in "
    "`widget/i18n-extra.js` — see Chapter 16. Placeholders such as `{family_name}` must match "
    "across every language for a key."
)
PAGEBREAK()

# ============================================================
# 5. CHECKPOINTER & PERSISTENCE
# ============================================================
H1("5. Checkpointer and analytics persistence")
FILEBOX("app/agent/checkpointer.py  ·  app/agent/setup_checkpointer.py  ·  app/persistence.py")

H2("The checkpointer — resumable graph state")
P(
    "After every node fires, LangGraph writes the full state snapshot to Postgres via "
    "`AsyncPostgresSaver`. That buys two things: **resumption** (a client re-posting the same "
    "`session_uuid` after a disconnect resumes at `current_node` with all slots intact) and an "
    "**audit** of every turn's state. `make_checkpointer()` is an async context manager scoped "
    "to a single run; the SSE router opens one per request."
)
P(
    "The checkpointer creates its own tables (`checkpoints`, `checkpoint_blobs`, "
    "`checkpoint_writes`, `checkpoint_migrations`) in the public schema. Alembic ignores them "
    "because they aren't in `Base.metadata`, so they coexist with the app's migrated tables. "
    "`app/agent/setup_checkpointer.py` is the one-time idempotent CLI (`python -m "
    "app.agent.setup_checkpointer`) that creates them; `deploy.sh` runs it on every deploy."
)
P(
    "Reminder on drivers: the checkpointer uses **psycopg**, the app uses **asyncpg**, both on "
    "the same database. That's why `config.py` exposes two URLs."
)

H2("app/persistence.py — the queryable mirror")
P(
    "The checkpoint blob resumes sessions but is not pleasant to query for analytics (funnel "
    "drop-offs, outcome counts, time-to-resolve). `persistence.py` keeps the shaped "
    "`conversations` and `messages` tables in lock-step with each turn. The SSE router is its "
    "only caller, because the router is the one place that sees both the user input and every "
    "event the graph emits — keeping these writes out of the nodes keeps the graph free of "
    "DB-write concerns."
)
P("The helpers:")
BULLET([
    "`upsert_conversation(session_uuid, …)` is idempotent on `session_uuid`: the first turn "
    "inserts, later turns update the timestamp / `current_node` / state. It carefully does not "
    "clobber `main_type_code` with NULL on a later turn that lacked a flow chip.",
    "`append_user_message` records the user's turn, tagging `content_type` by input kind "
    "(text / gate_choice / chip / picked_problem).",
    "`append_bot_text` and `append_bot_card` record each streamed event.",
    "`update_conversation_state` patches the row with the latest `current_node` and a "
    "**slimmed** state JSON (messages stripped — they live in their own table — and "
    "non-JSON-able values coerced).",
])
PAGEBREAK()

# ============================================================
# 6. NODES OVERVIEW
# ============================================================
H1("6. The nodes — overview and shared shape")
FILEBOX("app/agent/nodes/")
P(
    "Each file under `app/agent/nodes/` is one graph node exposing an `async def run(state) -> "
    "dict`. A node reads the state, optionally calls DeepSeek and the tools, and returns a "
    "partial state update. The next three chapters walk them by flow. First, the patterns they "
    "all share — learn these once and every node reads quickly."
)
BULLET([
    "**Structured output.** Narrow classification/extraction nodes call "
    "`get_chat_llm(temperature=0).with_structured_output(SomeModel)` so DeepSeek returns a "
    "validated pydantic object instead of free prose. Each node defines a small private "
    "`_Model` for exactly the fields it needs.",
    "**Partially deterministic gates.** Per the original backend plan, yes/no “gates” are "
    "not pure LLM calls. A button click arrives as a literal `\"yes\"`/`\"no\"` and bypasses the "
    "model entirely (the fast path); only free text is classified by the LLM; and a separate "
    "signal (a curated confidence score) can override a shaky “yes”. This makes the common "
    "path cheap and predictable.",
    "**Question handling.** Gates classify into yes / no / question / unclear. A **question** "
    "is answered (grounded on the candidate data via `fence`) and the gate is re-emitted; it "
    "does not count as a failed attempt. **Unclear** is re-asked once, then escalates to a "
    "human to avoid loops.",
    "**Cards vs. text.** Nodes append human prose to `messages` and structured UI payloads to "
    "`cards`. Card `kind`s (`product_results`, `recommendations`, `gate`, `problem_match`, "
    "`problem_candidates`, `custom_config`, `custom_config_form`, `outcome`, `suggest`) are a "
    "contract with the widget renderer — see Chapter 16.",
    "**Every string is translated.** User-facing prose goes through `t(key, lang)`; LLM prose "
    "is steered by the appended language instruction. Nodes read `state.get(\"language\")`.",
    "**current_node breadcrumbs.** Each return sets `current_node` to a dotted label "
    "(`guide.find.presented`) that lands in analytics and aids debugging.",
])
P(
    "The retired `echo.py` node is the Phase-2a placeholder kept as a plumbing smoke test; it "
    "just asks DeepSeek to echo the last message. It is wired only into `build_echo_graph`."
)
PAGEBREAK()

# ============================================================
# 7. ROUTING NODES
# ============================================================
H1("7. Routing nodes: entry_router and other.reclassify")
FILEBOX("app/agent/nodes/entry_router.py  ·  app/agent/nodes/other_reclassify.py")

H2("entry_router — first-touch classification")
P(
    "When a turn arrives without a `flow` (free-form text, not a chip click), this node picks "
    "exactly one of the four flows. It runs a deliberately narrow LLM classifier "
    "(`temperature=0`, structured `_Route`) that returns a flow code and a `customize` boolean "
    "— nothing else. No slot-filling, no chat; downstream nodes take over."
)
P(
    "There is also a **hybrid customise detector**. A short list of unambiguous phrases (“custom "
    "build”, “configure a…”, …) is matched in code; if any hit, the node forces "
    "`flow=guide` and the customise sub-flow regardless of the LLM, because missing that intent "
    "poisons routing (a stale `find_phase` would otherwise send the user to the gate instead of "
    "the configurator). When customise is detected it also wipes stale slots that could "
    "short-circuit dispatch. This belt-and-braces pattern — deterministic keyword check beside "
    "the LLM — recurs across the nodes whenever a miss is expensive."
)

H2("other.reclassify — second-chance routing")
P(
    "The `other` flow is the catch-all for greetings, off-topic chat, or input the first "
    "classifier couldn't read. Rather than dead-ending, this node makes one cheap LLM attempt "
    "to re-map the conversation into a real flow."
)
BULLET([
    "If the model is confident (≥ `CONFIDENCE_FLOOR` = 0.6) about presales/guide/postsales, "
    "it switches `state.flow` and the graph re-dispatches into that flow **in the same turn**.",
    "Otherwise it gives a friendly free-chat reply plus a `suggest` card of quick replies "
    "(“choosing a product” / “broken unit” / “talk to a human”).",
    "It does this at most `RECLASSIFY_MAX_ATTEMPTS` (2) times per session, then escalates to "
    "`outcome_human` so the user is never trapped in a loop.",
])
P(
    "Tuning these two constants changes how aggressively the bot tries to keep ambiguous users "
    "in a flow versus handing them to a human. They live at the top of the file."
)
PAGEBREAK()

# ============================================================
# 8. PRESALES & GUIDE NODES
# ============================================================
H1("8. Pre-sales and guide nodes")
FILEBOX("presales_figure_out.py  ·  guide_find.py  ·  guide_customize.py  ·  guide_happy_gate.py")

H2("presales.figure_out — application → family")
P(
    "Pre-sales is for users who haven't picked a product. The node runs in two phases. **Phase "
    "1** extracts the user's industry and application (structured `_Extraction`), and if either "
    "is missing asks one targeted follow-up. It also copies any literal model code the user "
    "typed into a `query` slot so the eventual handoff can land on that exact product. **Phase "
    "2** calls the `recommend_categories` tool, presents the top family with a `recommendations` "
    "card and a yes/no gate, and on “yes” hands off to `guide.find` by setting `flow=guide` "
    "and seeding `slots.filters.family`."
)
P(
    "A nice fallback: when the curated industry×application table has no match, the node gives "
    "the LLM the live catalog and asks it to pick the closest real family (validating the code "
    "isn't hallucinated) before escalating to a human. One cheap call beats a forced handoff "
    "when the seed simply doesn't cover that pair."
)

H2("guide.find — structured product search")
P(
    "Guide-find translates a free-form need into a `SearchProductsFilters` object and runs "
    "`search_products`. Its system prompt is built **from the live catalog** on each call "
    "(`_build_system_prompt` queries the active product types) so the list of valid family "
    "codes is never stale. Filters seeded by pre-sales merge with this turn's extraction, and "
    "any literal model code/SKU is carried in `query` so an exact-code search works even when "
    "the code doesn't map to a family."
)
P(
    "The most maintenance-relevant detail is the **relaxation ladder** for over-constrained "
    "searches: if the full filter set returns nothing, it retries with just the literal code, "
    "then with just the family. A shortlist beats “no results” every time. Results are "
    "presented as a `product_results` card plus a gate; the user's reply is then handled by "
    "`guide.happy_gate`."
)

H2("guide.customize — the configurator")
P(
    "When the user wants a build that doesn't match a stock SKU, this node walks them through "
    "the chosen family's `spec_schema`. It supports both a **structured form** (the widget "
    "submits `custom_modules`) and a **conversational** fallback (the LLM extracts module values "
    "from chat). Once at least `MIN_FILLED` (2) keys are set, it calls `build_custom_config`, "
    "emits a `custom_config` card plus the closest stock SKU, and re-uses `guide.happy_gate` for "
    "the quote decision. `MIN_FILLED` is intentionally low — the configurator is a discovery "
    "tool, not a CAD system; over-asking kills the flow."
)

H2("guide.happy_gate — interpreting the reply to candidates")
P(
    "Shared by find and customise. A chip click is a literal `yes`/`no`/`info_only`; free text "
    "is classified into yes / no / question / unclear. `info_only` (“I just needed the info”) "
    "resolves the conversation happily. A **question** triggers a second, grounded LLM call "
    "that answers using `slots.candidates` as the only source, then re-emits the gate — "
    "answering is the whole point of the gate, so it is not an escalation. Two unclear answers "
    "in a row route to a human. `gate_attempts` is bumped only on yes/no/unclear, never on a "
    "question turn."
)
PAGEBREAK()

# ============================================================
# 9. POST-SALES NODES
# ============================================================
H1("9. Post-sales nodes")
FILEBOX("postsales_identify.py  ·  postsales_match_kb.py  ·  postsales_fix_gate.py")

H2("postsales.identify — get a real symptom")
P(
    "Extracts an optional SKU and a **specific** symptom. The node is strict about what counts "
    "as a symptom: it must describe an observable behaviour (“overheating after 2 hours”), "
    "not a vague announcement (“my unit is broken”). Both the prompt and a code-level "
    "`_looks_vague` guard reject generic statements, because sending a vague symptom to the "
    "vector match yields a wrong answer. It does **not** block on SKU — many tickets land "
    "without one, and the family-agnostic match is still useful. Once a real symptom exists it "
    "sets `phase=ready` and the graph runs `match_kb`."
)
P(
    "There is a fast path: if the previous turn presented an ambiguous shortlist, the user's "
    "reply is treated directly as the refined symptom (the LLM tends to refuse to extract short "
    "noun-phrase labels), and the stale match state is cleared so `match_kb` re-runs cleanly."
)

H2("postsales.match_kb — vector match against the KB")
P(
    "Calls `find_problems(sku?, symptom_text)` and interprets the result against two thresholds:"
)
KV([
    ["`AMBIGUOUS_FLOOR` = 0.30", "Below the top similarity here, treat as no match at all and "
     "escalate to a human (offering it would mislead)."],
    ["`MATCH_FLOOR` = 0.55", "Below this but above the ambiguous floor, present a short pickable "
     "shortlist (`problem_candidates` card) instead of one confident answer."],
    ["≥ 0.55", "Present the single best problem and its top curated solution "
     "(`problem_match` card) plus a “Did that fix it?” gate."],
], headers=("Threshold", "Behaviour"))
P(
    "It also handles the deterministic path where the user clicked a candidate from a shortlist "
    "(`picked_problem_id`): it looks the problem up by id and presents it, skipping the vector "
    "search. Matched problems with no curated solution escalate gracefully. **Tuning these two "
    "floors** is the main lever for the precision/escalation trade-off in post-sales."
)

H2("postsales.fix_gate — did the fix work?")
P(
    "Three signals combine. A button click (`yes`/`no`) is strongest and bypasses the LLM. The "
    "curated solution's `confidence` (1–5) is the second: at or below `LOW_CONFIDENCE_FLOOR` "
    "(2), even a “yes” does **not** claim resolved — it escalates to a human to verify, "
    "because a false “resolved” is worse than a callback. Third is the LLM classification "
    "(yes / no / question / unclear) of free text, with one re-ask before falling through. As "
    "with the happy gate, a question is answered (grounded on the candidate solution) and the "
    "gate re-emitted. A “yes” ends at `outcome_resolved`; “no” (or the low-confidence "
    "case) ends at `outcome_human`."
)
PAGEBREAK()

# ============================================================
# 10. TERMINAL & POST-OUTCOME NODES
# ============================================================
H1("10. Terminal nodes and post-outcome chat")
FILEBOX("app/agent/nodes/outcomes.py  ·  app/agent/nodes/post_outcome_chat.py")

H2("outcomes.py — where flows end")
P(
    "Three terminal nodes set `outcome` and fire side effects, but they know nothing about CRM "
    "providers — they call `handle_sell` / `handle_human_handoff` / `handle_resolved` from "
    "`app.sideeffects` and surface the returned record ids on an `outcome` card."
)
BULLET([
    "`outcome_sell` — the user accepted a product. Builds a sales-qualified lead (RFQ), shows "
    "the chosen SKU, and surfaces the `rfq_id` / `division_code`.",
    "`outcome_human` — escalation. Picks a reason and priority from how the user got here "
    "(a failed fix is high priority; a rejected recommendation is normal) and creates a ticket.",
    "`outcome_resolved` — the curated fix worked. Logs a CRM activity (no record created) and "
    "offers a feedback prompt.",
])
P(
    "Because the provider is decoupled, swapping Zoho for Salesforce is an env change with no "
    "node edits — see Chapters 12–13."
)

H2("post_outcome_chat — staying polite after the end")
P(
    "Once a flow terminates, the structured conversation is done — a human will follow up "
    "out-of-band — but the user may still type. Dead-ending with “this conversation already "
    "wrapped up” is rude. This node handles any post-outcome turn with a single LLM call that "
    "reads the history plus the recorded outcome and replies conversationally. It is steered by "
    "tight rules (short replies, never re-open the flow, never promise timelines or prices, "
    "never tell the user the conversation is over) and an outcome-specific flavour line. It "
    "never changes state beyond appending its reply."
)
PAGEBREAK()

# ============================================================
# 11. TOOLS & SCHEMAS
# ============================================================
H1("11. Read-only tools and their schemas")
FILEBOX("app/tools/  ·  app/schemas/tools.py")
P(
    "Tools are plain async functions that take an `AsyncSession` and return pydantic objects. "
    "The agent calls them directly (no HTTP hop); the same functions are also exposed as thin "
    "FastAPI routes (Chapter 14) purely so they can be curl-tested without the agent. The "
    "request/response shapes live in `app/schemas/tools.py` so the JSON contract is identical "
    "on both call paths. The agent never sees raw DB rows or embedding vectors — only these "
    "structured results."
)

H2("The five tools")
KV([
    ["`search_products`", "Text + JSONB-filter SKU search. ILIKE on SKU/name/family/description "
     "plus typed filters (frame size, ratio, stages, torque, backlash, variant) against the "
     "per-product `specs` JSONB. Excludes inactive products."],
    ["`recommend_categories`", "Maps an (industry, application) pair to product families via the "
     "curated `use_case_product_types` table, ranked by `fit_score`. Exact match first, then a "
     "fuzzy ILIKE fallback for paraphrases."],
    ["`find_problems`", "Vector match a symptom against the problem KB. Resolves family from SKU "
     "(a strong prior that cuts cross-family false positives), embeds the symptom, and orders by "
     "pgvector cosine distance, attaching each problem's top solution."],
    ["`get_solution`", "Returns a problem plus all its solutions ordered by confidence. Used by "
     "the tool route; raises `ProblemNotFoundError` → 404."],
    ["`build_custom_config`", "Validates a family, echoes a normalised module payload, and picks "
     "the closest stock SKU by a tiny equality-count heuristic. No DB writes."],
], headers=("Tool", "What it does"))

H2("Notes that matter for maintenance")
BULLET([
    "**Stable signatures.** Phase 1 implemented `search_products` with ILIKE; the docstring "
    "notes Phase 3 could swap the text branch for a pgvector ANN query with the same signature, "
    "so agent code wouldn't change. Preserve signatures when you re-implement a tool.",
    "**Spec-key drift.** `search_products` handles stage count stored under `stages` / `stage` "
    "/ `stage_count`, and reads links from either the column or `specs`. These quirks exist "
    "because the catalog was machine-extracted; `app/seed/normalize.py` and `dedup_specs.py` "
    "fight the same drift (Chapter 15).",
    "**The cosine ordering footgun.** `find_problems` orders by the distance **expression**, "
    "not a label, because SQLAlchemy silently drops `order_by(\"dist\")` before a `LIMIT`, "
    "yielding an empty result set. If you touch that query, keep the expression ordering.",
    "**Thresholds live in the nodes, not the tools.** The tools return raw similarities; "
    "`match_kb` decides what counts as a match. Keep that separation.",
])
PAGEBREAK()

# ============================================================
# 12. EMBEDDINGS
# ============================================================
H1("12. The embeddings layer")
FILEBOX("app/embeddings.py")
P(
    "Embeddings power the post-sales symptom match. The module is a pluggable provider with "
    "three backends, chosen by `EMBEDDING_PROVIDER`, all emitting **1024-dim** vectors so the "
    "schema (`vector(1024)`) works regardless of choice."
)
KV([
    ["`hash`", "Deterministic SHA-based pseudo-embedding. Non-semantic — same text → same "
     "vector — so the pipeline works (uniqueness, dedup, index round-trips) with zero config. "
     "The CI/dev default. Vector matches are meaningless, so smoke tests skip semantic "
     "assertions under this provider."],
    ["`bge-m3`", "Local BGE-M3 via sentence-transformers. Production-quality and offline, but the "
     "first run downloads ~2 GB and CPU inference is slow on small hardware."],
    ["`dashscope`", "Alibaba Qwen text-embedding-v3 over HTTP. Low latency from China; needs "
     "`DASHSCOPE_API_KEY`. Batches at most 25 inputs per call."],
], headers=("Provider", "Notes"))
P(
    "`get_provider()` constructs the configured provider once (process-wide singleton); "
    "`embed_texts()` is the convenience wrapper the seed loader and `find_problems` call; "
    "`text_hash()` is a short SHA-1 used to skip re-embedding unchanged rows. **For production "
    "in China, set `dashscope`.** The `hash` default exists so CI never needs a key or a 2 GB "
    "download — if your symptom matches look random in dev, that's why: you're on `hash`."
)
PAGEBREAK()

# ============================================================
# 13. DATA MODEL
# ============================================================
H1("13. The database model")
FILEBOX("app/models/__init__.py  ·  content.py  ·  runtime.py  ·  sideeffects.py")
P(
    "Models are grouped by lifecycle. Importing the `app.models` package registers every table "
    "on `Base.metadata`, which is what Alembic autogenerate reads. The schema mirrors the "
    "original `DB_structure` PDF one-to-one."
)

H2("content.py — seeded, read-heavy")
KV([
    ["`main_conversation_types`", "The four routing branches (presales/guide/postsales/other)."],
    ["`use_cases`", "(industry, application) contexts; natural key for upsert."],
    ["`product_types`", "A product family. `spec_schema` (JSONB) declares which spec keys its "
     "SKUs expose, driving the configurator UI."],
    ["`use_case_product_types`", "Curated many-to-many fit between a use case and a family, with "
     "`fit_score` (1–5) and a rationale."],
    ["`products`", "A concrete SKU. `specs` (JSONB) holds family-specific keys. `status` gates "
     "visibility."],
    ["`problem_types`", "A known failure mode for a family, unique per (family, code)."],
    ["`solutions`", "Validated fixes per problem; `confidence` (1–5) feeds the fix gate."],
    ["`product_embeddings` / `problem_embeddings`", "pgvector(1024) rows with a `text_hash` so "
     "the loader skips unchanged re-embeds."],
], headers=("Table", "Role"))
P(
    "A reusable helper, `has_active_products()`, returns a scalar subquery used widely as "
    "`ProductType.id.in_(has_active_products())` so empty/inactive families never surface to "
    "the LLM or the user."
)

H2("runtime.py — written as users interact")
P(
    "`conversations` (one row per browser session, keyed by `session_uuid`, with a slim `state` "
    "JSON mirror, `outcome`, `ticket_id`, `rfq_id`, `division_code`, and timestamps) and "
    "`messages` (one row per turn: user, bot text, or bot card). These are the analytics tables "
    "`persistence.py` maintains."
)

H2("sideeffects.py — the outbound audit trail")
P(
    "`rfqs` and `tickets` are the agent's **internal** record of what was sent to the CRM "
    "(kept separately from CRM-side ids so an audit survives even if the external row is "
    "deleted). `crm_calls` and `email_calls` log every outbound provider call — the `log` "
    "providers write only here, and real providers add an audit row around each call. All four "
    "are provider-agnostic by design, so swapping Zoho for another CRM loses no history."
)
PAGEBREAK()

# ============================================================
# 14. MIGRATIONS & ROUTERS
# ============================================================
H1("14. Migrations and HTTP routers")

H2("alembic/ — schema migrations")
FILEBOX("alembic/env.py  ·  alembic/versions/*.py  ·  alembic.ini")
P(
    "`env.py` is async-aware and pulls the DB URL from `app.config` (so `alembic.ini` doesn't "
    "hardcode it) and imports `app.models` so autogenerate sees every table. The four migrations "
    "tell the build's history and are explained here once, by type, rather than line by line:"
)
KV([
    ["`c96b3d3b… phase 1 initial schema`", "All content + runtime tables."],
    ["`d4e2a9f5… phase 3 embeddings`", "`CREATE EXTENSION vector` + the two embedding tables. "
     "Uses **HNSW** indexes, not ivfflat: ivfflat needs training data and returns zero rows on "
     "the tiny seed sets; HNSW degrades gracefully."],
    ["`f2a8c91d… phase 4 side effects`", "rfqs, tickets, crm_calls, email_calls, plus "
     "conversations columns. Written to be idempotent on re-run."],
    ["`a7b1c2d3… add product_page_url`", "A one-column additive migration — the template for "
     "a routine schema tweak."],
])
P(
    "**To change the schema:** edit the model, run `alembic revision --autogenerate -m \"…\"`, "
    "review the generated file (autogenerate is a draft, not gospel — check JSONB defaults and "
    "server defaults), then `alembic upgrade head`. `deploy.sh` runs `upgrade head` on every "
    "deploy."
)

H2("app/routers/agent.py — the SSE endpoint")
FILEBOX("app/routers/agent.py")
P(
    "One endpoint, `POST /api/messages`, drives every interaction. The `MessageRequest` model is "
    "the full vocabulary the widget can send: `text`, `gate_choice`, `flow` (chip), `subflow` "
    "(customise), `picked_problem_id`, `language`, `custom_modules`, `contact_email`, and "
    "`force_human`. The handler sanitises input, persists the user turn, assembles the graph "
    "input (mapping each field into `messages`/`flow`/`slots`), then streams the graph with "
    "`graph.astream(…, stream_mode=\"updates\")`. For each node update it emits `bot_text` and "
    "`card` events (persisting them) and tracks the outcome; at the end it patches the "
    "conversation row and sends `done`. `_sse()` formats each event. `POST /api/sessions` is an "
    "optional helper that hands a fresh UUID to clients."
)

H2("app/routers/tools.py — curl-testable tool routes")
FILEBOX("app/routers/tools.py")
P(
    "Thin adapters over three tools (`search_products`, `recommend_categories`, "
    "`get_solution`) so they can be exercised over HTTP without an agent. They use the "
    "`get_session` dependency and the shared schemas. The agent does not go through these in "
    "production — it calls the underlying functions directly."
)
PAGEBREAK()

# ============================================================
# 15. SIDE EFFECTS / CRM / MAIL
# ============================================================
H1("15. Side effects: orchestration, CRM, and email")
FILEBOX("app/sideeffects/  ·  app/crm/  ·  app/mail/")
P(
    "This is the single seam between the agent's terminal nodes and the outside world. The "
    "design goal is that swapping a provider (Zoho → Salesforce, Aliyun → something else) is "
    "a one-file change and never touches a node."
)

H2("app/sideeffects/handlers.py — the orchestrator")
P("Each handler follows the same disciplined sequence:")
NUM([
    "Resolve a `Division` from the routing matrix.",
    "Build a provider-agnostic payload from agent state (`rfq.py`).",
    "Persist an `Rfq`/`Ticket` row **first** — the internal source of truth, written before any "
    "external call so it survives a provider outage.",
    "Call the configured CRM provider inside an audit wrapper that records a `crm_calls` row "
    "(ok or error) and updates the Rfq/Ticket with the returned record id.",
    "Send the division a notification email through the configured mail provider, recording an "
    "`email_calls` row.",
    "Update the parent `conversations` row with outcome / ids / division.",
])
P(
    "Errors are swallowed by default (`SIDEEFFECTS_SOFT_FAIL=true`) so a provider outage never "
    "breaks the user's terminal card; the audit rows let ops re-drive failures later. Flip it "
    "to `false` once production alerting exists."
)

H2("app/sideeffects/rfq.py — the only place that reads slot shape")
P(
    "`build_lead_payload` and `build_ticket_payload` are the **only** code that knows the "
    "structure of `slots`. CRM adapters take a `LeadPayload`/`TicketPayload` and never touch raw "
    "state, so the agent can evolve its slot vocabulary without breaking every integration. They "
    "assemble contact info, chosen SKU/family, notes, and a flattened transcript."
)

H2("app/sideeffects/routing.py + routing.yaml — lead routing")
P(
    "`routing.yaml` is a `(main_type, product_family) → division` matrix: a `divisions` map "
    "(code, name, inbox), a `families` map (family code → division), and `routes` with a "
    "fallback. `resolve_division` loads it once (cached) and resolves by precedence, returning a "
    "static default if the file is missing so dev still works. **Keep `families` in sync with "
    "`seed/product_types/`** — a new family with no entry falls through to Applications "
    "Engineering. (Note: the YAML references a `robotics` division for `kr_joint_series` that "
    "isn't defined in `divisions:`; such a miss silently falls back, so add the division if you "
    "want robot-reducer leads routed correctly.) The inbox addresses are placeholders and must "
    "be replaced with real distribution lists before go-live."
)

H2("app/crm/ — pluggable CRM")
P(
    "`__init__.py` defines the provider-neutral payload dataclasses (`LeadPayload`, "
    "`TicketPayload`, `ActivityPayload`, `CrmResult`, `TranscriptTurn`), the small `CrmProvider` "
    "Protocol (three operations), and the `get_provider()` factory keyed on `CRM_PROVIDER`. "
    "`log.py` is the default — synthetic ids, no third party, full audit. `zoho.py` implements "
    "the REST v6 API with the OAuth refresh-token grant, **region-aware hosts** (Kofon will use "
    "`cn`), in-process token caching with one 401 retry, and a `LeadPayload → Deal`, "
    "`TicketPayload → Case`, `ActivityPayload → Task` mapping. To add Salesforce, drop a "
    "`salesforce.py` implementing the same Protocol and wire it in the factory."
)

H2("app/mail/ — pluggable email")
P(
    "Same pattern: an `EmailMessage`/`EmailResult` pair, an `EmailProvider` Protocol, and a "
    "factory on `MAIL_PROVIDER`. `log.py` records would-be sends; `aliyun.py` implements Aliyun "
    "DirectMail with hand-rolled HMAC-SHA1 request signing (no SDK). The package is named `mail` "
    "to avoid colliding with the stdlib `email`."
)
PAGEBREAK()

# ============================================================
# 16. THE WIDGET
# ============================================================
H1("16. The browser widget")
FILEBOX("widget/widget.js  ·  api.js  ·  config.kofon.js  ·  i18n-extra.js  ·  index.html / index.zh.html  ·  widget.css")
P(
    "The widget is dependency-free vanilla JS, designed to drop into any site with four script "
    "tags. In this bundle it is served from `/` by the same FastAPI process."
)

H2("config.kofon.js — the site-specific layer")
P(
    "This is the file you edit to rebrand or reconfigure without touching widget internals. It "
    "holds branding (colours, agent name, avatar), the language list, the four primary "
    "**actions** (which map 1:1 to the four conversation types), the secondary **utilities** "
    "(custom build, lead times, datasheet, expo, talk-to-human), and quick links. Critically, "
    "`apiUrl` is the bare `\"/\"` because the widget and API share an origin — do not replace it "
    "with a fully-qualified URL or you re-introduce the CORS/mixed-content problems this bundle "
    "exists to avoid."
)

H2("api.js — the SSE client")
P(
    "Owns the `session_uuid` (generated client-side, stored in localStorage so a returning user "
    "resumes), and `streamMessage(apiUrl, payload, onEvent)`, which POSTs to `/api/messages` and "
    "parses the SSE stream, invoking `onEvent(name, data)` for each `bot_text`/`card`/`outcome` "
    "event and resolving on `done`. This is the mirror image of the server's `_sse()` format."
)

H2("widget.js — the UI and card renderer")
P(
    "The ~1800-line core: it builds the chat DOM, manages open/minimised state and the language "
    "switcher, renders the welcome screen from `config`, and — most importantly for backend "
    "engineers — **renders one method per card kind**. The dispatch maps each `kind` the nodes "
    "emit to a renderer:"
)
CODE(
    "product_results    → _renderProductResultsCard\n"
    "recommendations    → _renderRecommendationsCard\n"
    "gate               → _renderGateCard       (Yes / No / dismiss chips)\n"
    "outcome            → _renderOutcomeCard\n"
    "problem_candidates → _renderProblemCandidatesCard\n"
    "problem_match      → _renderProblemMatchCard\n"
    "custom_config_form → _renderCustomConfigFormCard"
)
P(
    "Gate chips post `{gate_choice: \"yes\"|\"no\"|\"info_only\"}`; a picked candidate posts "
    "`{picked_problem_id}`; the config form posts `{custom_modules}`. **The card `kind`s and "
    "payload fields are a contract between the nodes and these renderers** — if you add a card "
    "kind in a node, add a renderer here, and vice-versa. widget.js carries an EN-only "
    "translation table; all other languages load from i18n-extra.js before `mount()`."
)

H2("i18n-extra.js, index.html / index.zh.html, widget.css")
BULLET([
    "`i18n-extra.js` — the DE/FR/RU/JA/KO/ZH translations for the **widget's** own strings "
    "(separate from the backend's `app/i18n.py`). Adding a language touches both tables.",
    "`index.html` / `index.zh.html` — the demo marketing page that hosts the widget (and a "
    "Chinese variant). Most of it is a static Kofon landing page; the four `<script>` tags at "
    "the bottom mount the widget.",
    "`widget.css` — all widget styling, self-contained so the host page can't bleed in.",
])
PAGEBREAK()

# ============================================================
# 17. SEED DATA & LOADERS
# ============================================================
H1("17. Seed content and the loaders")
FILEBOX("seed/  ·  app/seed/")
P(
    "`seed/` is the source of truth for everything the bot reads: editable YAML/CSV that domain "
    "experts maintain. `app/seed/` is the code that loads it into Postgres and the scripts that "
    "help create it. The loader is idempotent — running it twice changes nothing the second "
    "time — and upserts by natural key."
)

H2("The seed file types (one explanation each, not per file)")
P(
    "There are dozens of per-family files; they all follow one of these shapes. Edit the YAML/CSV "
    "and re-run `python -m app.seed.load`."
)
KV([
    ["`main_conversation_types.yaml`", "The four router branches. Natural key `code`. Rarely "
     "edited."],
    ["`use_cases.yaml`", "Industry×application contexts. Natural key (industry, application)."],
    ["`product_types/<code>.yaml`", "**One product family per file.** Declares `code`, `name`, "
     "`family`, `description`, `product_page_url`, and the `spec_schema` (the typed spec keys "
     "its SKUs expose, with labels/enums that drive the configurator). Example: "
     "`caesarplanetary.yaml`."],
    ["`products/<code>.yaml`", "**The SKUs in one family.** A `product_type_code` link plus a "
     "list of products, each with `sku`, `name`, a `specs` map keyed by the family's "
     "spec_schema, datasheet/CAD URLs, lead time, and status."],
    ["`problems/<code>.yaml`", "Known problems for a family, each with nested `solutions` "
     "(summary, markdown body, `confidence`, optional SOP/RMA URLs, `escalate_if`)."],
    ["`use_case_fits.csv`", "The many-to-many (industry, application, family) fit rows with "
     "`fit_score` and rationale. The CSV is authoritative — rows removed from it are deleted on "
     "load."],
], headers=("File type", "What it holds"))

H2("app/seed/load.py — the loader")
P(
    "Loads in FK order: conversation types → use cases → product types → products → fits "
    "→ problems+solutions, then refreshes embeddings. `_upsert` does bulk `ON CONFLICT DO "
    "UPDATE`, taking care to fill the **union** of keys across heterogeneous rows (otherwise "
    "SQLAlchemy compiles the INSERT from the first row's keys and silently drops columns later "
    "rows add). Solutions have no natural key, so each problem's solutions are deleted and "
    "reinserted on every run. The loader validates references (an unknown `product_type_code` "
    "raises a clear error naming the file), which makes most seed mistakes fail loudly at load "
    "time."
)

H2("app/seed/embed.py — content → vectors")
P(
    "Builds a short, consistent text chunk per product type / product / problem, hashes it, and "
    "only re-embeds rows whose hash changed (dropping stale rows in the same pass). Run "
    "automatically at the end of `load`, or standalone with `python -m app.seed.embed`. Keep "
    "the chunk shapes short and technical so embeddings cluster cleanly."
)

H2("The catalog-ingest scripts (data prep)")
BULLET([
    "`extract_url.py` / `extract_pdf.py` — point at a Kofon product page or datasheet PDF; they "
    "fetch it, clean it (BeautifulSoup / pdfplumber, rendering PDF tables as markdown so column "
    "structure survives), and ask DeepSeek (structured `ProductTypeDraft`) to emit draft YAML "
    "under `seed/drafts/`. Drafts are **not** auto-loaded — a human reviews and promotes them "
    "into `seed/product_types/` + `seed/products/`. This is how new families get added at scale.",
    "`normalize.py` — one-time, comment-preserving cleanup that lifts datasheet/CAD links out of "
    "`specs` onto the product columns and unifies frame size to a clean integer `frame_size_mm` "
    "plus an optional string label, so the numeric search filter actually matches. Idempotent; "
    "`--check` dry-runs.",
    "`dedup_specs.py` — collapses synonym spec keys (e.g. `stage`/`stages`, "
    "`input_speed_rpm_max`/`max_input_speed_rpm`) that the draft-merge step left behind, so the "
    "configurator doesn't render duplicate fields. Per-family merge rules live in a `MERGES` "
    "dict at the top.",
])
P(
    "`seed/README.md` documents the file table and natural keys. The `extract_*`, `normalize`, "
    "and `dedup_specs` scripts are run by maintainers, not at deploy time."
)
PAGEBREAK()

# ============================================================
# 18. DEPLOYMENT
# ============================================================
H1("18. Deployment and project configuration")
FILEBOX("deploy.sh  ·  docker-compose.yml  ·  .env.example  ·  pyproject.toml")

H2("deploy.sh — the one command")
P(
    "Idempotent and safe to re-run after a code update. It checks prerequisites, copies "
    "`.env.example` → `.env` if missing (and warns if `DEEPSEEK_API_KEY` looks unset), brings "
    "up Postgres via docker compose and waits for the health check, creates a venv and installs "
    "the package editable, then runs the three provisioning steps in order: `alembic upgrade "
    "head`, `python -m app.agent.setup_checkpointer`, `python -m app.seed.load`. Finally it "
    "starts the server with `python -m app.serve`. Flags: `--no-run` (provision only), "
    "`--reset-db` (wipe the Postgres volume first — destructive)."
)
CODE(
    "./deploy.sh                # provision + run on :8001\n"
    "./deploy.sh --no-run       # provision only\n"
    "./deploy.sh --reset-db     # wipe DB volume, then provision + run\n\n"
    "# public URL for a demo, in another terminal:\n"
    "cloudflared tunnel --url http://localhost:8001"
)

H2("docker-compose.yml — local Postgres")
P(
    "One service: `pgvector/pgvector:pg16` (the pgvector extension preinstalled), with a named "
    "volume for persistence and a `pg_isready` health check `deploy.sh` waits on. `docker "
    "compose down -v` wipes the data."
)

H2(".env.example — the configuration surface")
P(
    "The annotated template for `.env`. It groups Postgres, DeepSeek, embeddings, CRM, and email "
    "settings, each with inline guidance on when a block is required (e.g. the Zoho block only "
    "when `CRM_PROVIDER=zoho`). The defaults are the demo-safe combination: `EMBEDDING_PROVIDER="
    "hash`, `CRM_PROVIDER=log`, `MAIL_PROVIDER=log`, `SIDEEFFECTS_SOFT_FAIL=true`. **For "
    "production you must at minimum set a real `DEEPSEEK_API_KEY`, switch the embedding provider "
    "to `dashscope` (or `bge-m3`), and configure Zoho + Aliyun** with the `cn` region."
)

H2("pyproject.toml — dependencies and tooling")
P(
    "Declares the runtime dependencies (FastAPI, SQLAlchemy[asyncio]+asyncpg, Alembic, the "
    "LangChain/LangGraph + DeepSeek stack, the psycopg-based checkpointer, pgvector, httpx, "
    "BeautifulSoup, pdfplumber) with comments explaining each Phase's additions, plus a `dev` "
    "extra (ruff, pytest, ruamel.yaml). Ruff line length is 100; pytest runs in asyncio auto "
    "mode. Only the `app*` package is installed."
)
PAGEBREAK()

# ============================================================
# 19. SMOKE TESTS
# ============================================================
H1("19. Smoke tests")
FILEBOX("app/agent/smoke_*.py")
P(
    "The `smoke_2*` / `smoke_3*` / `smoke_4.py` scripts are runnable end-to-end checks for each "
    "Phase, kept in the tree rather than a `tests/` folder because they exercise the real graph "
    "against a seeded database. They are explained here as a **type**, not individually: each "
    "installs the Windows loop policy, drives one flow or tool path, and asserts on the result. "
    "For example, `smoke_3a.py` confirms the embeddings provider returns 1024-dim vectors, that "
    "`app.seed.embed` populated the tables, and that `find_problems` ranks the planted "
    "`backlash_exceeds_spec` problem first for a matching symptom (skipping the semantic "
    "assertion when the non-semantic `hash` provider is active). Run them after `deploy.sh` to "
    "confirm a deployment is wired correctly; run the one matching the area you changed as a "
    "fast regression check."
)
PAGEBREAK()

# ============================================================
# 20. MAINTENANCE COOKBOOK
# ============================================================
H1("20. Maintenance cookbook")
P(
    "Common tasks, each pointing back to the detailed chapters. After any content or schema "
    "change, re-run the relevant provisioning step and the matching smoke test."
)

H3("Add or update a product family")
NUM([
    "Either hand-write `seed/product_types/<code>.yaml` + `seed/products/<code>.yaml`, or run "
    "`python -m app.seed.extract_url <url>` / `extract_pdf <pdf>` and review the draft.",
    "Run `python -m app.seed.normalize` and `python -m app.seed.dedup_specs` if the data was "
    "machine-extracted.",
    "Add the family code to `app/sideeffects/routing.yaml` under `families:` so leads route to "
    "the right division.",
    "Add `use_case_fits.csv` rows if the family should surface in pre-sales recommendations.",
    "`python -m app.seed.load` (re-embeds automatically). Verify with `smoke_3a`/`smoke_2*`.",
])

H3("Add a known problem and fix")
NUM([
    "Add a problem block (with nested solutions and a `confidence`) to "
    "`seed/problems/<family>.yaml`. Remember `confidence ≤ 2` will force a human verify even "
    "on a “yes” in the fix gate.",
    "`python -m app.seed.load`. The new problem text is embedded automatically.",
])

H3("Add a language")
NUM([
    "Backend: add the code to `SUPPORTED_LANGUAGES` / `LANGUAGE_NAMES` and a column to every "
    "entry in `_T` in `app/i18n.py` (EN is the fallback, so partial is OK).",
    "Widget: add the language to `config.kofon.js` `languages` and the strings to "
    "`widget/i18n-extra.js`.",
])

H3("Swap the CRM or email provider")
NUM([
    "For an existing adapter, just set `CRM_PROVIDER=zoho` (or `MAIL_PROVIDER=aliyun`) and the "
    "required credentials in `.env`. No code changes.",
    "For a new provider, add `app/crm/<name>.py` (or `app/mail/<name>.py`) implementing the "
    "Protocol, and register it in the factory in the package `__init__.py`.",
])

H3("Tune the bot's behaviour")
KV([
    ["Match strictness (post-sales)", "`MATCH_FLOOR` / `AMBIGUOUS_FLOOR` in `postsales_match_kb.py`."],
    ["Low-confidence escalation", "`LOW_CONFIDENCE_FLOOR` in `postsales_fix_gate.py`."],
    ["Reclassify aggressiveness", "`CONFIDENCE_FLOOR` / `RECLASSIFY_MAX_ATTEMPTS` in `other_reclassify.py`."],
    ["Configurator minimum inputs", "`MIN_FILLED` in `guide_customize.py`."],
    ["Which LLM a node uses", "`get_chat_llm(model=…)` / `get_reasoner_llm()` in `app/agent/llm.py`."],
    ["Lead routing", "`app/sideeffects/routing.yaml`."],
], headers=("To change…", "Edit…"))

H3("Add a graph node")
NUM([
    "Create `app/agent/nodes/<name>.py` with `async def run(state) -> dict`, using "
    "`system_message()` for any prompt and `t()` for any fixed string.",
    "Register it in `build_graph` (`app/agent/graph.py`) and extend the dispatch function plus "
    "the allowed-target list of every conditional edge that can reach or leave it.",
    "If it emits a new card kind, add a renderer in `widget/widget.js`.",
])

P("")
P(
    "When in doubt, follow the data: a user turn enters at `app/routers/agent.py`, flows through "
    "`app/agent/graph.py` into one node, which calls a tool in `app/tools/` and reads content "
    "seeded from `seed/`. Every layer in that path has its chapter above.",
)

# ---------------- save ----------------
# Drop a live Table of Contents on its own page between the cover and Ch.0.
add_toc_before_first_heading(doc, title="Contents")
add_page_numbers(doc)
out = "Kofon_Chatbot_Engineering_Manual.docx"
doc.save(out)
print(f"wrote {out}  ({len(doc.paragraphs)} paragraphs)")
